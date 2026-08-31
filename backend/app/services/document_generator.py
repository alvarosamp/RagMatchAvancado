from __future__ import annotations

import hashlib
import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable
from zoneinfo import ZoneInfo

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.table import Table
from docx.oxml.ns import qn

from app.services.proposal_generator import DEFAULT_COMPANY, build_notice_proposal_docx


TEMPLATE_ROOT = Path(__file__).resolve().parents[1] / "templates"
DOCUMENT_TEMPLATE_ROOT = TEMPLATE_ROOT / "documents"
LETTERHEAD_BACKGROUND_A4 = DOCUMENT_TEMPLATE_ROOT / "papel_timbrado_fundo_a4.png"
LETTERHEAD_BACKGROUND_LETTER = DOCUMENT_TEMPLATE_ROOT / "papel_timbrado_fundo_carta.png"


def _today_brasilia() -> str:
    return datetime.now(ZoneInfo("America/Sao_Paulo")).strftime("%d/%m/%Y")


@dataclass(frozen=True)
class DocumentTemplate:
    id: str
    name: str
    template_path: Path
    required_fields: tuple[str, ...]
    generator: Callable[[Any, dict[str, Any], dict[str, Any]], bytes]

    @property
    def version(self) -> str:
        digest = hashlib.sha256(self.template_path.read_bytes())
        digest.update(LETTERHEAD_BACKGROUND_A4.read_bytes())
        digest.update(LETTERHEAD_BACKGROUND_LETTER.read_bytes())
        return digest.hexdigest()[:12]


def _commercial(notice: Any, company: dict[str, Any], options: dict[str, Any]) -> bytes:
    return build_notice_proposal_docx(notice, company=company, options=options)


REGISTRY: dict[str, DocumentTemplate] = {
    "commercial_proposal": DocumentTemplate(
        "commercial_proposal", "Proposta comercial", TEMPLATE_ROOT / "proposals" / "modelo_proposta.docx",
        ("company.razao_social", "company.cnpj"), _commercial,
    ),
    "unified_declaration": DocumentTemplate(
        "unified_declaration", "Declaração Unificada", DOCUMENT_TEMPLATE_ROOT / "declaracao_unificada.docx",
        ("company.razao_social", "company.cnpj", "company.endereco", "company.representante", "company.cpf_representante", "notice.number", "notice.organ"),
        lambda notice, company, options: _build_unified(notice, company, options),
    ),
    "feasibility_declaration": DocumentTemplate(
        "feasibility_declaration", "Declaração de Exequibilidade", DOCUMENT_TEMPLATE_ROOT / "declaracao_exequibilidade.docx",
        ("company.razao_social", "company.cnpj", "company.endereco", "company.representante", "company.cpf_representante", "company.rg_representante", "notice.number", "notice.organ"),
        lambda notice, company, options: _build_feasibility(notice, company, options),
    ),
}


def list_templates() -> list[dict[str, Any]]:
    return [
        {"id": item.id, "name": item.name, "template_version": item.version, "required_fields": list(item.required_fields)}
        for item in REGISTRY.values()
    ]


def generation_preview(notice: Any, template_id: str, company: dict[str, Any], options: dict[str, Any]) -> dict[str, Any]:
    template = _template(template_id)
    company_fields = _normalize_company(notice, company, options)
    fields = _preview_fields(notice, company_fields, options)
    missing = [
        field for field in template.required_fields
        if not field.startswith("company.") and not _field_value(field, fields)
    ]
    if template_id in {"commercial_proposal", "feasibility_declaration"}:
        for product in getattr(notice, "notice_products", []) or []:
            if getattr(product, "selected_for_dispute", True) is False:
                continue
            if getattr(product, "quantity", None) is None:
                missing.append(f"items.{getattr(product, 'item_number', product.id)}.quantity")
            if template_id == "commercial_proposal" and getattr(product, "unit_price", None) is None:
                missing.append(f"items.{getattr(product, 'item_number', product.id)}.unit_price")
    return {"template": {"id": template.id, "name": template.name, "template_version": template.version}, "fields": fields, "missing_fields": sorted(set(missing))}


def generate_document(notice: Any, template_id: str, company: dict[str, Any], options: dict[str, Any]) -> tuple[bytes, dict[str, Any]]:
    preview = generation_preview(notice, template_id, company, options)
    if preview["missing_fields"] and not options.get("allow_missing"):
        raise ValueError("Dados pendentes: " + ", ".join(preview["missing_fields"]))
    template = _template(template_id)
    content = template.generator(notice, preview["fields"]["company"], options)
    document = Document(BytesIO(content))
    apply_letterhead(document)
    output = BytesIO()
    document.save(output)
    validated = output.getvalue()
    _validate_docx(validated)
    return validated, preview


def generated_filename(notice: Any, template_id: str) -> str:
    number = re.sub(r"[^A-Za-z0-9._-]+", "_", str(getattr(notice, "number", None) or getattr(notice, "id", "edital"))).strip("_")
    return f"{template_id}_{number or 'edital'}.docx"


def _template(template_id: str) -> DocumentTemplate:
    template = REGISTRY.get(template_id)
    if template is None:
        raise ValueError("Modelo de documento invalido.")
    return template


def _preview_fields(notice: Any, company: dict[str, Any], options: dict[str, Any]) -> dict[str, Any]:
    organ = getattr(getattr(notice, "organ", None), "name", None)
    return {
        "company": company,
        "notice": {
            "id": getattr(notice, "id", None), "number": getattr(notice, "number", None),
            "bid_number": getattr(notice, "bid_number", None), "organ": organ,
            "municipality": getattr(notice, "municipality_name", None), "modality": getattr(notice, "modality", None),
        },
        "signer": options.get("signer") or {},
        "emission_city": options.get("city") or company.get("cidade") or getattr(notice, "municipality_name", None),
        "emission_date": options.get("date") or _today_brasilia(),
    }


def _normalize_company(notice: Any, company: dict[str, Any] | None, options: dict[str, Any] | None) -> dict[str, Any]:
    supplied = company or {}
    opts = options or {}
    signer = opts.get("signer") or {}
    city = opts.get("city") or supplied.get("cidade") or DEFAULT_COMPANY.get("cidade") or DEFAULT_COMPANY.get("cidade_uf") or getattr(notice, "municipality_name", None) or ""
    defaults = {
        **DEFAULT_COMPANY,
        "cidade": city,
        "nacionalidade": "brasileiro(a)",
        "estado_civil": "",
        "funcao_representante": signer.get("role") or "representante legal",
    }
    merged = {**defaults, **{key: value for key, value in supplied.items() if value not in (None, "")}}
    if not merged.get("representante"):
        merged["representante"] = signer.get("name") or ""
    if not merged.get("funcao_representante"):
        merged["funcao_representante"] = signer.get("role") or "representante legal"
    if not merged.get("cidade"):
        merged["cidade"] = city
    return merged


def _field_value(path: str, fields: dict[str, Any]) -> Any:
    current: Any = fields
    for part in path.split("."):
        if not isinstance(current, dict):
            return None
        current = current.get(part)
    return current


def _build_unified(notice: Any, company: dict[str, Any], options: dict[str, Any]) -> bytes:
    document = Document(str(DOCUMENT_TEMPLATE_ROOT / "declaracao_unificada.docx"))
    signer = options.get("signer") or {}
    representative = company.get("representante") or signer.get("name") or ""
    cpf = company.get("cpf_representante") or signer.get("cpf") or ""
    nationality = company.get("nacionalidade") or ""
    marital = company.get("estado_civil") or ""
    qualifications = ", ".join(value for value in (nationality, marital) if value)
    organ = getattr(getattr(notice, "organ", None), "name", None) or ""
    intro = (
        f"A empresa {company.get('razao_social', '')}, inscrita no CNPJ sob o Nº {company.get('cnpj', '')}, "
        f"estabelecida em {company.get('endereco', '')}, através do seu representante legal {representative}"
        f"{', ' + qualifications if qualifications else ''}, inscrito no CPF sob o nº {cpf}. DECLARA, que:"
    )
    _set_paragraph(document.paragraphs[1], intro)
    notice_number = getattr(notice, "bid_number", None) or getattr(notice, "number", None) or ""
    for paragraph in document.paragraphs:
        if "Edital do Pregão Eletrônico" in paragraph.text:
            _set_paragraph(paragraph, re.sub(r"Edital do Pregão Eletrônico[^,]*", f"Edital do Pregão Eletrônico nº {notice_number}, promovido por {organ}", paragraph.text, count=1))
    city = options.get("city") or company.get("cidade") or getattr(notice, "municipality_name", None) or ""
    emission_date = options.get("date") or _today_brasilia()
    _set_paragraph(document.paragraphs[27], f"{city}, {emission_date}.")
    _set_paragraph(document.paragraphs[30], f"{signer.get('name') or representative}\nCPF: {signer.get('cpf') or cpf}")
    return _save(document)


def _build_feasibility(notice: Any, company: dict[str, Any], options: dict[str, Any]) -> bytes:
    document = Document(str(DOCUMENT_TEMPLATE_ROOT / "declaracao_exequibilidade.docx"))
    signer = options.get("signer") or {}
    products = [item for item in getattr(notice, "notice_products", []) or [] if getattr(item, "selected_for_dispute", True) is not False]
    item_numbers = ", ".join(str(getattr(item, "item_number", "")) for item in products)
    organ = getattr(getattr(notice, "organ", None), "name", None) or ""
    process_number = options.get("process_number") or getattr(notice, "number", None) or ""
    auction_number = getattr(notice, "bid_number", None) or getattr(notice, "number", None) or ""
    intro = (
        f"A empresa {company.get('razao_social', '')}, inscrita no CNPJ sob o nº {company.get('cnpj', '')}, "
        f"com sede em {company.get('endereco', '')}, por intermédio de seu representante legal "
        f"{company.get('representante') or signer.get('name') or ''}, portador(a) da Carteira de Identidade nº "
        f"{company.get('rg_representante', '')} e CPF nº {company.get('cpf_representante') or signer.get('cpf') or ''}, "
        f"DECLARA que possui capacidade de honrar o fornecimento dos itens {item_numbers} do processo licitatório nº "
        f"{process_number} de {organ}, {getattr(notice, 'modality', None) or 'Pregão Eletrônico'} nº {auction_number}. "
        "Conforme planilha de custos abaixo:"
    )
    _set_paragraph(document.paragraphs[1], intro)
    _fill_feasibility_table(document.tables[0], products, options)
    justification = options.get("justification") or ""
    _set_paragraph(document.paragraphs[3], f"Justificativa: {justification}")
    city = options.get("city") or company.get("cidade") or getattr(notice, "municipality_name", None) or ""
    emission_date = options.get("date") or _today_brasilia()
    _set_paragraph(document.paragraphs[6], f"{city}, {emission_date}.")
    _set_paragraph(document.paragraphs[8], signer.get("name") or company.get("representante") or "")
    _set_paragraph(document.paragraphs[9], signer.get("role") or company.get("funcao_representante") or "")
    _set_paragraph(document.paragraphs[10], company.get("razao_social") or "")
    return _save(document)


def _fill_feasibility_table(table: Table, products: list[Any], options: dict[str, Any]) -> None:
    sample = deepcopy(table.rows[-1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    item_values = options.get("item_values") or {}
    for product in products:
        table._tbl.append(deepcopy(sample))
        row = table.rows[-1]
        key = str(getattr(product, "id", None) or getattr(product, "item_number", ""))
        supplied = item_values.get(key) or item_values.get(str(getattr(product, "item_number", ""))) or {}
        catalog = getattr(product, "catalog_product", None)
        cost = supplied.get("cost", getattr(product, "cost", None))
        quantity = getattr(product, "quantity", None)
        total = supplied.get("initial_max_total", getattr(product, "reference_total_price", None))
        if total is None and getattr(product, "reference_price", None) is not None and quantity is not None:
            total = float(product.reference_price) * float(quantity)
        values = [
            getattr(product, "item_number", ""), getattr(product, "description", ""),
            supplied.get("brand", getattr(catalog, "brand", None)), supplied.get("model", getattr(catalog, "model", None)),
            quantity, getattr(product, "unit", None), _money_or_blank(cost), supplied.get("net_markup"),
            _money_or_blank(supplied.get("freight")), _money_or_blank(supplied.get("taxes")), _money_or_blank(total),
        ]
        for cell, value in zip(row.cells, values):
            _set_cell(cell, value)


def apply_letterhead(target: Document, letterhead: Document | None = None) -> None:
    """Apply one full-page branded background behind the document content.

    ``letterhead`` remains accepted for compatibility with older callers, but
    the generated documents use the precomposed A4/Letter background assets.
    """
    for section in target.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.header_distance = 0
        section.footer_distance = 0
        _clear_header_footer(section.header.part)
        _clear_header_footer(section.footer.part)

        page_width = int(section.page_width)
        page_height = int(section.page_height)
        background_path = (
            LETTERHEAD_BACKGROUND_A4
            if page_height / page_width > 1.35
            else LETTERHEAD_BACKGROUND_LETTER
        )
        paragraph = section.header.paragraphs[0]
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1
        shape = paragraph.add_run().add_picture(
            str(background_path),
            width=section.page_width,
            height=section.page_height,
        )
        _inline_picture_to_page_background(shape._inline, page_width, page_height)


def _clear_header_footer(part: Any) -> None:
    for child in list(part.element):
        part.element.remove(child)
    part.element.append(OxmlElement("w:p"))
    for rel_id, rel in list(part.rels.items()):
        if rel.reltype in {RT.IMAGE, RT.HYPERLINK}:
            part.drop_rel(rel_id)


def _inline_picture_to_page_background(inline: Any, page_width: int, page_height: int) -> None:
    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "0", "behindDoc": "1",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)
    for axis in ("H", "V"):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "0"
        position.append(offset)
        anchor.append(position)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(page_width))
    extent.set("cy", str(page_height))
    anchor.append(extent)
    effect_extent = inline.find(qn("wp:effectExtent"))
    if effect_extent is not None:
        anchor.append(deepcopy(effect_extent))
    anchor.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(deepcopy(element))

    for element in anchor.iter(qn("a:off")):
        element.set("x", "0")
        element.set("y", "0")
    for element in anchor.iter(qn("a:ext")):
        element.set("cx", str(page_width))
        element.set("cy", str(page_height))
    inline.getparent().replace(inline, anchor)


def _validate_docx(content: bytes) -> None:
    reopened = Document(BytesIO(content))
    for index, section in enumerate(reopened.sections, start=1):
        header_images = sum(1 for rel in section.header.part.rels.values() if rel.reltype == RT.IMAGE)
        footer_images = sum(1 for rel in section.footer.part.rels.values() if rel.reltype == RT.IMAGE)
        anchors = list(section.header._element.iter(qn("wp:anchor")))
        if header_images != 1 or footer_images != 0 or len(anchors) != 1 or anchors[0].get("behindDoc") != "1":
            raise ValueError(f"O fundo do papel timbrado ficou invalido na secao {index}.")


def _set_paragraph(paragraph: Any, value: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    (paragraph.runs[0] if paragraph.runs else paragraph.add_run()).text = str(value or "")


def _set_cell(cell: Any, value: Any) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    _set_paragraph(paragraph, "" if value is None else str(value))
    for extra in cell.paragraphs[1:]:
        _set_paragraph(extra, "")


def _money_or_blank(value: Any) -> str:
    if value in (None, ""):
        return ""
    return "R$ " + f"{float(value):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _save(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()
