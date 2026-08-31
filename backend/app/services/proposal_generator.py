from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import _Cell, _Row, Table

TEMPLATE_PATH = (
    Path(__file__).resolve().parents[1]
    / "templates"
    / "proposals"
    / "modelo_proposta.docx"
)


DEFAULT_COMPANY = {
    "razao_social": "TOR TECNOLOGIA E INDUSTRIA LTDA",
    "cnpj": "63.799.679/0001-40",
    "endereco": "RUA COMENDADOR CUSTODIO RIBEIRO, N 137, SALA 16",
    "bairro": "CENTRO",
    "cidade_uf": "SANTA RITA DO SAPUCAI/MG",
    "cep": "37.536-072",
    "telefone": "(31) 99898-9039",
    "email": "governo@tor.tec.br",
    "representante": "ALLAN CESAR DE PAIVA BARREIROS",
    "cpf_representante": "052.832.226-59",
    "rg_representante": "MG 10179237",
    "banco": "BANCO INTER, Agencia: 0001-9, Conta Corrente: 49933223-7",
}


def build_notice_proposal_docx(
    notice: Any,
    *,
    company: dict[str, Any] | None = None,
    options: dict[str, Any] | None = None,
) -> bytes:
    """Generate a commercial proposal DOCX.

    Items won by our company are preferred. If no item result exists yet, the
    proposal is generated as a commercial preview with active linked items.
    """
    company_data = {**DEFAULT_COMPANY, **(company or {})}
    options = options or {}
    won_items = _collect_won_items(notice) or _collect_preview_items(notice)
    if not won_items:
        raise ValueError("Nao ha itens vinculados para gerar proposta.")

    document = Document(str(TEMPLATE_PATH))
    _fill_header_table(document.tables[0], notice, company_data, options)
    _fill_items_table(document.tables[1], won_items)
    _replace_common_text(document, notice, company_data, options, won_items)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def proposal_filename(notice: Any) -> str:
    number = _safe_filename(getattr(notice, "number", None) or getattr(notice, "tor_id", None))
    if not number:
        number = _safe_filename(getattr(notice, "id", "edital"))
    return f"proposta_{number}.docx"


def _collect_won_items(notice: Any) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    products_by_id = {product.id: product for product in getattr(notice, "notice_products", [])}
    for result in getattr(notice, "notice_item_results", []) or []:
        winner_type = getattr(result, "winner_type", None)
        winner_value = getattr(winner_type, "value", winner_type)
        if winner_value != "us":
            continue
        product = products_by_id.get(result.notice_product_id) or result.notice_product
        if product is None:
            continue
        if getattr(product, "selected_for_dispute", True) is False:
            continue
        quantity = _required_number("quantidade", getattr(result, "winning_quantity", None), getattr(product, "quantity", None))
        unit_price = _required_number(
            "preco comercial",
            getattr(result, "winning_price", None),
            getattr(product, "unit_price", None),
        )
        catalog = getattr(product, "catalog_product", None)
        brand = (
            getattr(result, "winner_brand", None)
            or getattr(catalog, "brand", None)
            or ""
        )
        model = (
            getattr(result, "winner_model", None)
            or getattr(catalog, "model", None)
            or ""
        )
        items.append(
            {
                "item": getattr(product, "item_number", None) or str(len(items) + 1),
                "description": getattr(product, "description", None) or getattr(catalog, "name", None) or "Produto",
                "unit": getattr(product, "unit", None) or getattr(catalog, "unit", None) or "UN",
                "quantity": quantity,
                "unit_price": unit_price,
                "brand_model": " / ".join(part for part in [brand, model] if part) or "A DEFINIR",
                "total": quantity * unit_price,
                "delivery_deadline": getattr(product, "delivery_deadline", None),
                "warranty": getattr(product, "warranty", None),
            }
        )
    return sorted(items, key=lambda item: _sort_key(item["item"]))


def _collect_preview_items(notice: Any) -> list[dict[str, Any]]:
    inactive_result_ids = {
        getattr(result, "notice_product_id", None)
        for result in getattr(notice, "notice_item_results", []) or []
        if getattr(getattr(result, "winner_type", None), "value", getattr(result, "winner_type", None)) != "us"
    }
    items: list[dict[str, Any]] = []
    for product in getattr(notice, "notice_products", []) or []:
        if getattr(product, "id", None) in inactive_result_ids:
            continue
        if getattr(product, "selected_for_dispute", True) is False:
            continue
        catalog = getattr(product, "catalog_product", None)
        unit_price = _required_number("preco comercial", getattr(product, "unit_price", None))
        quantity = _required_number("quantidade", getattr(product, "quantity", None))
        brand = getattr(catalog, "brand", None) or ""
        model = getattr(catalog, "model", None) or ""
        items.append(
            {
                "item": getattr(product, "item_number", None) or str(len(items) + 1),
                "description": getattr(product, "description", None) or getattr(catalog, "name", None) or "Produto",
                "unit": getattr(product, "unit", None) or getattr(catalog, "unit", None) or "UN",
                "quantity": quantity,
                "unit_price": unit_price,
                "brand_model": " / ".join(part for part in [brand, model] if part) or "A DEFINIR",
                "total": quantity * unit_price,
                "delivery_deadline": getattr(product, "delivery_deadline", None),
                "warranty": getattr(product, "warranty", None),
            }
        )
    return sorted(items, key=lambda item: _sort_key(item["item"]))


def _fill_header_table(
    table: Table,
    notice: Any,
    company: dict[str, Any],
    options: dict[str, Any],
) -> None:
    organ = getattr(notice, "organ", None)
    portal = getattr(notice, "portal", None)
    organ_name = getattr(organ, "name", None) or getattr(notice, "municipality_name", None) or "ORGAO"
    process_number = (
        options.get("process_number")
        or getattr(notice, "bid_number", None)
        or getattr(notice, "number", None)
        or getattr(notice, "tor_id", None)
        or ""
    )
    modality = options.get("auction_number") or getattr(notice, "modality", None) or "PREGAO ELETRONICO"
    judgment = options.get("judgment_type") or getattr(notice, "bi_criterion", None) or "conforme edital"
    portal_name = getattr(portal, "name", None)
    if portal_name:
        organ_name = f"{organ_name}\nPORTAL: {portal_name}"

    rows = [
        [organ_name] * 5,
        [
            f"PROCESSO N {process_number}",
            f"PROCESSO N {process_number}",
            modality,
            modality,
            modality,
        ],
        [""] * 5,
        [f"RAZAO SOCIAL: {company['razao_social']}"] * 4 + [f"CNPJ/CPF: {company['cnpj']}"],
        [f"ENDERECO: {company['endereco']}"] * 3 + [f"BAIRRO: {company['bairro']}"] * 2,
        [
            f"CIDADE/UF: {company['cidade_uf']}",
            f"CEP: {company['cep']}",
            f"CEP: {company['cep']}",
            f"CEP: {company['cep']}",
            f"TELEFONE: {company['telefone']}",
        ],
        [f"REPRESENTANTE LEGAL: {company['representante']}"] * 4 + [f"CPF: {company['cpf_representante']}"],
        [f"RG: {company['rg_representante']}", f"E-mail: {company['email']}"] * 2 + [f"E-mail: {company['email']}"],
    ]
    for row_index, row_values in enumerate(rows):
        if row_index >= len(table.rows):
            break
        if row_index == 2:
            _set_judgment_row(table.rows[row_index], judgment)
            continue
        for cell, value in zip(table.rows[row_index].cells, row_values):
            _set_cell_text_preserving_style(cell, value)


def _set_judgment_row(row: _Row, judgment: Any) -> None:
    cells = row.cells
    if not cells:
        return
    target = cells[0]
    if len(cells) > 1:
        target = target.merge(cells[-1])
    _set_cell_text_preserving_style(target, f"TIPO DE JULGAMENTO: {judgment}")


def _fill_items_table(table: Table, items: list[dict[str, Any]]) -> None:
    if len(table.rows) < 3:
        raise ValueError("Template de proposta sem tabela de itens valida.")

    sample_row = deepcopy(table.rows[1]._tr)
    total_row = deepcopy(table.rows[-1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    total = 0.0
    for item in items:
        table._tbl.append(deepcopy(sample_row))
        row = table.rows[-1]
        total += item["total"]
        values = [
            str(item["item"]),
            item["description"],
            item["unit"],
            _format_quantity(item["quantity"]),
            _money(item["unit_price"]),
            item["brand_model"],
            _money(item["total"]),
        ]
        for cell, value in zip(row.cells, values):
            _set_cell_text_preserving_style(cell, value)

    table._tbl.append(total_row)
    total_text = _money(total)
    total_in_words = valor_por_extenso(total)
    for index, cell in enumerate(table.rows[-1].cells):
        _set_cell_text_preserving_style(
            cell,
            "VALOR TOTAL"
            if index < len(table.rows[-1].cells) - 1
            else f"{total_text}\n({total_in_words})",
        )


def _replace_common_text(
    document: Document,
    notice: Any,
    company: dict[str, Any],
    options: dict[str, Any],
    won_items: list[dict[str, Any]],
) -> None:
    current_city = options.get("proposal_city") or "Santa Rita do Sapucai"
    current_date = options.get("proposal_date") or _date_pt_br(datetime.now())
    delivery_term = options.get("delivery_term") or _delivery_term_from_items(won_items)
    validity_term = _proposal_validity_term(notice, options)
    warranty_term = options.get("warranty") or _term_from_items(
        won_items,
        "warranty",
        singular_fallback="conforme garantia prevista no edital",
        multiple_fallback="conforme garantias previstas no edital por item",
    )
    replacements = {
        "O prazo de validade da proposta": (
            f"O prazo de validade da proposta e de {validity_term}."
        ),
        "Nome do banco indicado para o pagamento": (
            f"Nome do banco indicado para o pagamento: {company['banco']}."
        ),
        "Prazo de entrega/execu": (
            f"Prazo de entrega/execucao: "
            f"{delivery_term}, "
            "contados do recebimento da Solicitacao de Fornecimento/Ordem de Servicos."
        ),
        "Prazo de Garantia": (
            f"Prazo de Garantia: {warranty_term}, "
            "contra defeito de fabricacao, contados a partir da data da entrega."
        ),
        "Santa Rita do Sapuca": f"{current_city}, {current_date}.",
        "Allan Cesar de Paiva Barreiros": (
            f"{company['representante']}\n CPF: {company['cpf_representante']}"
        ),
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        for needle, replacement in replacements.items():
            if needle in text:
                _set_paragraph_text(paragraph, replacement)
                break


def _delivery_term_from_items(items: list[dict[str, Any]]) -> str:
    return _term_from_items(
        items,
        "delivery_deadline",
        singular_fallback="conforme prazo de entrega previsto no edital",
        multiple_fallback="conforme prazos de entrega previstos no edital por item",
    )


def _term_from_items(
    items: list[dict[str, Any]],
    key: str,
    *,
    singular_fallback: str,
    multiple_fallback: str,
) -> str:
    deadlines: list[str] = []
    for item in items:
        value = str(item.get(key) or "").strip()
        if value and value not in deadlines:
            deadlines.append(value)

    if not deadlines:
        return singular_fallback
    if len(deadlines) == 1:
        return deadlines[0]

    return multiple_fallback


def _proposal_validity_term(notice: Any, options: dict[str, Any]) -> str:
    explicit_days = options.get("validity_days")
    if explicit_days:
        return f"{explicit_days} dias"

    validity = str(getattr(notice, "proposal_validity", None) or "").strip()
    if validity:
        return validity

    return "conforme prazo de validade previsto no edital"


def _set_cell_text(cell: _Cell, text: Any) -> None:
    cell.text = str(text or "")


def _set_cell_text_preserving_style(cell: _Cell, text: Any) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = str(text or "")
        return
    _set_paragraph_text(paragraphs[0], str(text or ""))
    for paragraph in paragraphs[1:]:
        _set_paragraph_text(paragraph, "")


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _number(*values: Any) -> float:
    for value in values:
        if value in (None, ""):
            continue
        try:
            return float(value)
        except (TypeError, ValueError):
            continue
    return 0.0


def _required_number(label: str, *values: Any) -> float:
    for value in values:
        if value not in (None, ""):
            return float(value)
    raise ValueError(f"Preencha {label} dos itens antes de gerar o documento.")


def _money(value: float) -> str:
    text = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {text}"


def valor_por_extenso(value: float) -> str:
    reais = int(value)
    centavos = int(round((value - reais) * 100))
    parts: list[str] = []
    if reais:
        parts.append(_int_to_words(reais) + (" real" if reais == 1 else " reais"))
    if centavos:
        suffix = " centavo" if centavos == 1 else " centavos"
        parts.append(_int_to_words(centavos) + suffix)
    return " e ".join(parts) if parts else "zero reais"


def _int_to_words(number: int) -> str:
    if number == 0:
        return "zero"

    scales = [
        (1_000_000_000, "bilhao", "bilhoes"),
        (1_000_000, "milhao", "milhoes"),
        (1_000, "mil", "mil"),
    ]
    parts: list[str] = []
    remaining = number
    for scale, singular, plural in scales:
        chunk = remaining // scale
        if chunk == 0:
            continue
        remaining %= scale
        if scale == 1_000 and chunk == 1:
            parts.append("mil")
        else:
            label = singular if chunk == 1 else plural
            parts.append(f"{_int_to_words_under_1000(chunk)} {label}")

    if remaining:
        if parts and remaining < 100:
            parts.append("e " + _int_to_words_under_1000(remaining))
        else:
            parts.append(_int_to_words_under_1000(remaining))
    return ", ".join(parts).replace(", e ", " e ").replace("mil, ", "mil ")


def _int_to_words_under_1000(number: int) -> str:
    units = [
        "",
        "um",
        "dois",
        "tres",
        "quatro",
        "cinco",
        "seis",
        "sete",
        "oito",
        "nove",
    ]
    teens = [
        "dez",
        "onze",
        "doze",
        "treze",
        "quatorze",
        "quinze",
        "dezesseis",
        "dezessete",
        "dezoito",
        "dezenove",
    ]
    tens = [
        "",
        "",
        "vinte",
        "trinta",
        "quarenta",
        "cinquenta",
        "sessenta",
        "setenta",
        "oitenta",
        "noventa",
    ]
    hundreds = [
        "",
        "cento",
        "duzentos",
        "trezentos",
        "quatrocentos",
        "quinhentos",
        "seiscentos",
        "setecentos",
        "oitocentos",
        "novecentos",
    ]
    if number == 100:
        return "cem"
    if number < 10:
        return units[number]
    if number < 20:
        return teens[number - 10]
    if number < 100:
        ten, unit = divmod(number, 10)
        return tens[ten] if unit == 0 else f"{tens[ten]} e {units[unit]}"

    hundred, rest = divmod(number, 100)
    return hundreds[hundred] if rest == 0 else f"{hundreds[hundred]} e {_int_to_words_under_1000(rest)}"


def _format_quantity(value: float) -> str:
    if float(value).is_integer():
        return str(int(value))
    return str(value).replace(".", ",")


def _date_pt_br(value: datetime) -> str:
    months = [
        "janeiro",
        "fevereiro",
        "marco",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    return f"{value.day:02d} de {months[value.month - 1]} de {value.year}"


def _sort_key(value: Any) -> tuple[int, str]:
    text = str(value or "")
    try:
        return (0, f"{int(text):08d}")
    except ValueError:
        return (1, text)


def _safe_filename(value: Any) -> str:
    text = "".join(ch if ch.isalnum() else "_" for ch in str(value or ""))
    return "_".join(part for part in text.split("_") if part)[:80]
