from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.services.document_generator import generate_document, generation_preview, list_templates
from app.services.proposal_generator import build_notice_proposal_docx


def _fixture():
    catalog = SimpleNamespace(name="Switch", brand="TOR", model="SW24", unit="UN")
    product = SimpleNamespace(
        id="p1", item_number="1", description="Switch gerenciavel", quantity=2, unit="UN",
        cost=100.0, unit_price=180.0, reference_price=200.0, reference_total_price=400.0,
        selected_for_dispute=True, catalog_product=catalog,
    )
    notice = SimpleNamespace(
        id="n1", number="90001/2026", bid_number="90001/2026", municipality_name="Sao Paulo",
        modality="Pregao Eletronico", organ=SimpleNamespace(name="Municipio de Sao Paulo"),
        portal=SimpleNamespace(name="Compras.gov.br"), notice_products=[product], notice_item_results=[],
    )
    company = {
        "razao_social": "EMPRESA TESTE LTDA", "cnpj": "12.345.678/0001-90",
        "endereco": "Rua Teste, 123, Sao Paulo/SP", "representante": "Maria da Silva",
        "cpf_representante": "123.456.789-00", "rg_representante": "MG-12.345.678", "cidade": "Sao Paulo",
    }
    options = {"justification": "Compra em volume.", "signer": {"name": "Maria da Silva", "cpf": "123.456.789-00", "role": "Diretora"}}
    return notice, company, options


def test_registry_exposes_three_extensible_templates():
    assert {item["id"] for item in list_templates()} == {
        "commercial_proposal", "unified_declaration", "feasibility_declaration"
    }


@pytest.mark.parametrize("template_id", ["commercial_proposal", "unified_declaration", "feasibility_declaration"])
def test_generated_documents_use_one_full_page_background_behind_content(template_id):
    notice, company, options = _fixture()
    content, _ = generate_document(notice, template_id, company, options)
    document = Document(BytesIO(content))
    for section in document.sections:
        assert sum(rel.reltype == RT.IMAGE for rel in section.header.part.rels.values()) == 1
        assert sum(rel.reltype == RT.IMAGE for rel in section.footer.part.rels.values()) == 0
        anchors = list(section.header._element.iter("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"))
        assert len(anchors) == 1
        assert anchors[0].get("behindDoc") == "1"
        assert anchors[0].find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH").get("relativeFrom") == "page"
        assert anchors[0].find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV").get("relativeFrom") == "page"


def test_declarations_replace_examples_and_create_one_item_row():
    notice, company, options = _fixture()
    unified, _ = generate_document(notice, "unified_declaration", company, options)
    unified_doc = Document(BytesIO(unified))
    unified_text = "\n".join(paragraph.text for paragraph in unified_doc.paragraphs)
    assert "Pregão Eletrônico nº 018/2026" not in unified_text
    assert "XX de julho de 2026" not in unified_text
    assert "EMPRESA TESTE LTDA" in unified_text
    feasibility, _ = generate_document(notice, "feasibility_declaration", company, options)
    feasibility_doc = Document(BytesIO(feasibility))
    assert len(feasibility_doc.tables[0].rows) == 2
    table_text = "\n".join(cell.text for row in feasibility_doc.tables[0].rows for cell in row.cells)
    assert "106" not in table_text and "107" not in table_text
    assert "Switch gerenciavel" in table_text


def test_reference_price_is_not_used_as_commercial_price():
    notice, company, _ = _fixture()
    notice.notice_products[0].unit_price = None
    with pytest.raises(ValueError, match="preço comercial"):
        build_notice_proposal_docx(notice, company=company)


def test_proposal_keeps_judgment_type_in_one_field():
    notice, company, options = _fixture()
    content = build_notice_proposal_docx(notice, company=company, options=options)
    document = Document(BytesIO(content))
    judgment_texts = {
        cell.text.strip()
        for cell in document.tables[0].rows[2].cells
        if cell.text.strip()
    }
    assert judgment_texts == {"TIPO DE JULGAMENTO: conforme edital"}


def test_preview_reports_missing_quantity_without_turning_it_into_zero():
    notice, company, options = _fixture()
    notice.notice_products[0].quantity = None
    preview = generation_preview(notice, "feasibility_declaration", company, options)
    assert "items.1.quantity" in preview["missing_fields"]
