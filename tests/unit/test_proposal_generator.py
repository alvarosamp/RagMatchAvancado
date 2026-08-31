from types import SimpleNamespace

from docx import Document

from app.services.proposal_generator import build_notice_proposal_docx


def _enum_value(value: str):
    return SimpleNamespace(value=value)


def test_build_notice_proposal_includes_all_won_items(tmp_path):
    products = [
        SimpleNamespace(
            id="item-1",
            item_number="1",
            description="Switch gerenciavel 24 portas",
            unit="UN",
            quantity=2,
            unit_price=None,
            reference_price=500,
            catalog_product=SimpleNamespace(
                name="Switch 24 portas",
                brand="TP-Link",
                model="SG2428",
                unit="UN",
            ),
        ),
        SimpleNamespace(
            id="item-2",
            item_number="2",
            description="Roteador corporativo",
            unit="UN",
            quantity=1,
            unit_price=None,
            reference_price=900,
            catalog_product=SimpleNamespace(
                name="Roteador",
                brand="MikroTik",
                model="RB4011",
                unit="UN",
            ),
        ),
        SimpleNamespace(
            id="item-3",
            item_number="3",
            description="Item perdido",
            unit="UN",
            quantity=1,
            unit_price=None,
            reference_price=100,
            catalog_product=None,
        ),
    ]
    notice = SimpleNamespace(
        id="notice-1",
        number="123/2026",
        tor_id="TOR-001",
        modality="Pregao Eletronico 001/2026",
        organ=SimpleNamespace(name="Municipio de Teste"),
        portal=SimpleNamespace(name="Portal Teste"),
        notice_products=products,
        notice_item_results=[
            SimpleNamespace(
                notice_product_id="item-1",
                notice_product=products[0],
                winner_type=_enum_value("us"),
                winning_quantity=2,
                winning_price=450,
                winner_brand="TP-Link",
                winner_model="SG2428",
            ),
            SimpleNamespace(
                notice_product_id="item-2",
                notice_product=products[1],
                winner_type="us",
                winning_quantity=1,
                winning_price=850,
                winner_brand="MikroTik",
                winner_model="RB4011",
            ),
            SimpleNamespace(
                notice_product_id="item-3",
                notice_product=products[2],
                winner_type="competitor",
                winning_quantity=1,
                winning_price=90,
                winner_brand=None,
                winner_model=None,
            ),
        ],
    )

    content = build_notice_proposal_docx(notice)
    path = tmp_path / "proposta.docx"
    path.write_bytes(content)

    document = Document(str(path))
    item_table_text = "\n".join(
        cell.text for row in document.tables[1].rows for cell in row.cells
    )

    assert "Switch gerenciavel 24 portas" in item_table_text
    assert "Roteador corporativo" in item_table_text
    assert "Item perdido" not in item_table_text
    assert "R$ 900,00" in item_table_text
    assert "R$ 850,00" in item_table_text
    assert "R$ 1.750,00" in item_table_text
    assert "mil setecentos e cinquenta reais" in item_table_text


def test_build_notice_proposal_uses_notice_delivery_deadline(tmp_path):
    product = SimpleNamespace(
        id="item-1",
        item_number="1",
        description="Access point corporativo",
        unit="UN",
        quantity=3,
        unit_price=None,
        reference_price=700,
        warranty="36 (trinta e seis) meses",
        delivery_deadline="30 (trinta) dias corridos",
        catalog_product=SimpleNamespace(
            name="Access point",
            brand="Ruckus",
            model="R650",
            unit="UN",
        ),
    )
    notice = SimpleNamespace(
        id="notice-2",
        number="TOR-CRM-456",
        tor_id="TOR-002",
        bid_number="Pregao Eletronico 456/2026",
        modality="Pregao Eletronico 002/2026",
        bi_criterion="Menor preco por item",
        proposal_validity="60 (sessenta) dias",
        organ=SimpleNamespace(name="Municipio de Prazo"),
        portal=SimpleNamespace(name="Portal Teste"),
        notice_products=[product],
        notice_item_results=[
            SimpleNamespace(
                notice_product_id="item-1",
                notice_product=product,
                winner_type="us",
                winning_quantity=3,
                winning_price=650,
                winner_brand="Ruckus",
                winner_model="R650",
            ),
        ],
    )

    content = build_notice_proposal_docx(notice)
    path = tmp_path / "proposta_prazo.docx"
    path.write_bytes(content)

    document = Document(str(path))
    proposal_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    header_text = "\n".join(
        cell.text for row in document.tables[0].rows for cell in row.cells
    )

    assert "Prazo de entrega/execução: 30 (trinta) dias corridos" in proposal_text
    assert "O prazo de validade da proposta é de 60 (sessenta) dias." in proposal_text
    assert "Prazo de Garantia: 36 (trinta e seis) meses" in proposal_text
    assert "05 (cinco) dias uteis" not in proposal_text
    assert "90 dias" not in proposal_text
    assert "12 meses" not in proposal_text
    assert "PROCESSO Nº Pregao Eletronico 456/2026" in header_text
    assert "Menor preco por item" in header_text


def test_build_notice_proposal_preview_before_won_items(tmp_path):
    product = SimpleNamespace(
        id="item-1",
        item_number="1",
        description="Kit switch com transceiver",
        unit="UN",
        quantity=2,
        unit_price=1200,
        reference_price=1500,
        warranty=None,
        delivery_deadline=None,
        catalog_product=SimpleNamespace(
            name="Switch",
            brand="Huawei",
            model="S5735",
            unit="UN",
            min_price=1200,
        ),
    )
    notice = SimpleNamespace(
        id="notice-preview",
        number="PREV-001",
        tor_id="TOR-PREV",
        modality="Pregao Eletronico",
        organ=SimpleNamespace(name="Municipio Preview"),
        portal=None,
        notice_products=[product],
        notice_item_results=[],
    )

    content = build_notice_proposal_docx(notice)
    path = tmp_path / "proposta_previa.docx"
    path.write_bytes(content)

    document = Document(str(path))
    item_table_text = "\n".join(
        cell.text for row in document.tables[1].rows for cell in row.cells
    )

    assert "Kit switch com transceiver" in item_table_text
    assert "Huawei / S5735" in item_table_text
    assert "R$ 2.400,00" in item_table_text


def test_build_notice_proposal_preview_ignores_inactive_items(tmp_path):
    active = SimpleNamespace(
        id="item-active",
        item_number="1",
        description="Item ativo",
        unit="UN",
        quantity=1,
        unit_price=100,
        reference_price=120,
        selected_for_dispute=True,
        catalog_product=None,
    )
    out_of_dispute = SimpleNamespace(
        id="item-out",
        item_number="2",
        description="Item fora da disputa",
        unit="UN",
        quantity=1,
        unit_price=200,
        reference_price=220,
        selected_for_dispute=False,
        catalog_product=None,
    )
    lost = SimpleNamespace(
        id="item-lost",
        item_number="3",
        description="Item perdido no preview",
        unit="UN",
        quantity=1,
        unit_price=300,
        reference_price=330,
        selected_for_dispute=True,
        catalog_product=None,
    )
    notice = SimpleNamespace(
        id="notice-preview-filter",
        number="PREV-002",
        tor_id="TOR-PREV-2",
        modality="Pregao Eletronico",
        organ=None,
        portal=None,
        notice_products=[active, out_of_dispute, lost],
        notice_item_results=[
            SimpleNamespace(
                notice_product_id="item-lost",
                notice_product=lost,
                winner_type="competitor",
            ),
        ],
    )

    content = build_notice_proposal_docx(notice)
    path = tmp_path / "proposta_previa_filtrada.docx"
    path.write_bytes(content)

    document = Document(str(path))
    item_table_text = "\n".join(
        cell.text for row in document.tables[1].rows for cell in row.cells
    )

    assert "Item ativo" in item_table_text
    assert "Item fora da disputa" not in item_table_text
    assert "Item perdido no preview" not in item_table_text
