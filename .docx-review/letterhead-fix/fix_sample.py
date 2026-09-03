from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\vish8\OneDrive\Documentos\RagMatchAvan-ado")
sys.path.insert(0, str(ROOT / "backend"))

from app.services.document_generator import _validate_docx, apply_letterhead  # noqa: E402


SOURCE = Path(r"C:\Users\vish8\Downloads\commercial_proposal_SUSPENSO_-_2026_08_17_6.docx")
OUTPUT = ROOT / "output" / "commercial_proposal_corrigida.docx"
REPLACEMENTS = {
    "PROCESSO N ": "PROCESSO Nº ",
    "RAZAO SOCIAL:": "RAZÃO SOCIAL:",
    "ENDERECO:": "ENDEREÇO:",
    "O prazo de validade da proposta e de": "O prazo de validade da proposta é de",
    "Prazo de entrega/execucao:": "Prazo de entrega/execução:",
    "Solicitacao de Fornecimento": "Solicitação de Fornecimento",
    "Ordem de Servicos": "Ordem de Serviços",
    "Agencia:": "Agência:",
    "milhoes": "milhões",
    "milhao": "milhão",
    "bilhoes": "bilhões",
    "bilhao": "bilhão",
    "Santa Rita do Sapucai": "Santa Rita do Sapucaí",
}


def update_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        text = run.text
        for old, new in REPLACEMENTS.items():
            text = text.replace(old, new)
        run.text = text


document = Document(str(SOURCE))
for paragraph in document.paragraphs:
    update_paragraph(paragraph)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                update_paragraph(paragraph)

apply_letterhead(document)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
buffer = BytesIO()
document.save(buffer)
content = buffer.getvalue()
_validate_docx(content)
OUTPUT.write_bytes(content)
print(OUTPUT)
